from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
import uuid
import re
from difflib import SequenceMatcher
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime, timezone
from io import BytesIO
from urllib.parse import quote

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from lesson_content import LESSONS, LESSON_INDEX, GRADE_NAMES

BASE_RESOURCES = ["الكتاب المدرسي", "الحاسوب", "بروكسيما (جهاز العرض)"]


def _nums(*idx):
    return "-".join(str(i) for i in idx)


def _all(n):
    return _nums(*range(1, n + 1))


VARIANTS = [
    {"variant": 1, "name": "الحوار والمناقشة والتعلم التعاوني", "tagline": "النمط التقليدي المطوّر — عرض عملي ومناقشة موجّهة",
     "strategies": lambda n: [("الحوار والمناقشة", _all(n)), ("العرض العملي (التوضيح)", _all(n)), ("التعلم التعاوني", _nums(*range(1, min(n, 2) + 1))),
                              ("التعلم الذاتي", _nums(*range(max(1, n - 1), n + 1))), ("العصف الذهني", "1"), ("حل المشكلات", str(n))],
     "exec": "({i}) من خلال طريقة الحوار والمناقشة والعرض العملي على جهاز العرض يوضّح المعلم ما يأتي: {pt}، ثم يناقش الطلاب في مجموعات ثنائية وينفّذون الخطوات على أجهزتهم بحيث يحقق الطالب الهدف: {obj}، ويُعزَّز ذلك بـ{act}.",
     "extra_res": [], "extra_form": "أسئلة شفهية مباشرة أثناء العرض العملي", "hw": "{hw}"},
    {"variant": 2, "name": "الاستقصاء والاكتشاف", "tagline": "الطالب يكتشف المفاهيم بنفسه عبر أسئلة استقصائية وتجريب",
     "strategies": lambda n: [("الاستقصاء والاكتشاف", _all(n)), ("العصف الذهني", "1"), ("الحوار والمناقشة", _all(n)), ("التعلم بالأقران", _nums(*range(2, n + 1))),
                              ("خرائط المفاهيم", _all(n)), ("التعلم بالممارسة", _nums(*range(2, n + 1)))],
     "exec": "({i}) يطرح المعلم سؤالاً استقصائياً مرتبطاً بالهدف ({obj})، ويستكشف الطلاب الإجابة بأنفسهم من الكتاب المدرسي وبالتجريب على الجهاز، ثم يستنتجون ما يأتي: {pt}، ويسجّلون استنتاجاتهم في خريطة مفاهيم جماعية، ويتحقّق المعلم من صحة الاستنتاج عبر {act}.",
     "extra_res": ["ورقة عمل استقصائية", "خريطة مفاهيم فارغة"], "extra_form": "ملاحظة سجل الاستنتاجات وخرائط المفاهيم", "hw": "{hw}، مع كتابة سؤال استقصائي جديد حول الدرس لمناقشته في الحصة القادمة"},
    {"variant": 3, "name": "التعلم التعاوني بمجموعات الخبراء", "tagline": "كل مجموعة تتخصص في جزء من الدرس ثم تعلّمه للآخرين",
     "strategies": lambda n: [("التعلم التعاوني (مجموعات الخبراء)", _all(n)), ("التدريس بالأقران", _all(n)), ("العرض والتقديم", _all(n)),
                              ("الحوار والمناقشة", _all(n)), ("لعب الأدوار", _nums(*range(1, min(n, 2) + 1))), ("التقييم بالأقران", str(n))],
     "exec": "({i}) تُكلَّف مجموعة الخبراء رقم ({i}) بدراسة ما يأتي من الكتاب المدرسي: {pt}، وتجربتها عملياً، ثم يعود كل خبير إلى مجموعته الأصلية ليشرح لزملائه ويدرّبهم بحيث يحقق كل طالب الهدف: {obj}، ويعرض ممثل المجموعة الملخص أمام الصف، ويُقيَّم الفهم بـ{act}.",
     "extra_res": ["بطاقات مجموعات الخبراء", "قائمة شطب للتقييم بالأقران"], "extra_form": "تقييم الأقران بقائمة شطب لعرض كل مجموعة", "hw": "{hw}، وإعداد ملخص خبير (5 نقاط) عن الجزء الذي درسته المجموعة"},
    {"variant": 4, "name": "التعلم القائم على المهام العملية والمشروع", "tagline": "مهام تطبيقية متدرجة تنتهي بمنتج رقمي للطالب",
     "strategies": lambda n: [("التعلم القائم على المهام / المشروع", _all(n)), ("التعلم بالممارسة (العملي)", _all(n)), ("حل المشكلات", _nums(*range(2, n + 1))),
                              ("التعلم التعاوني", _all(n)), ("النمذجة والعرض العملي", "1"), ("التغذية الراجعة الفورية", _all(n))],
     "exec": "({i}) المهمة العملية ({i}): ينفّذ الطلاب على أجهزتهم مهمة تطبيقية تتطلب أن {obj}، مستندين إلى ما يأتي: {pt}، ويتابع المعلم الأداء ويقدّم التغذية الراجعة الفورية، ويُوثَّق الناتج ضمن منتج الدرس النهائي مع الربط بـ{act}.",
     "extra_res": ["قائمة رصد للمهمة العملية", "نموذج المنتج النهائي"], "extra_form": "قائمة رصد لأداء المهام العملية على الجهاز", "hw": "إكمال المنتج العملي للدرس وتسليمه إلكترونياً، إضافة إلى {hw}"},
    {"variant": 5, "name": "الصف المعكوس والتعلم الذاتي", "tagline": "اطّلاع مسبق في المنزل ثم تطبيق ومسابقة إلكترونية في الحصة",
     "strategies": lambda n: [("الصف المعكوس", _all(n)), ("التعلم الذاتي", _all(n)), ("فكّر – زاوج – شارك", _all(n)), ("الحوار والمناقشة", _all(n)),
                              ("التعلم الإلكتروني والألعاب التعليمية", _all(n)), ("التقييم الذاتي", str(n))],
     "exec": "({i}) يطالع الطلاب قبل الحصة ما يأتي من الكتاب المدرسي (تعلّم ذاتي): {pt}، وفي الحصة يُطبَّق نمط فكّر–زاوج–شارك للإجابة عن أسئلة حول الهدف ({obj})، ثم يختبرون فهمهم بمسابقة إلكترونية قصيرة عبر منصة الاختبارات، ويُكمل الطلاب {act} فردياً مع دعم المعلم للمتعثرين.",
     "extra_res": ["منصة الاختبارات الإلكترونية (مسابقة قصيرة)", "ورقة تهيئة منزلية"], "extra_form": "نتائج المسابقة الإلكترونية القصيرة عبر المنصة", "hw": "قراءة الدرس التالي من الكتاب والإجابة عن أسئلة ورقة التهيئة، إضافة إلى {hw}"},
]


def build_variant(lesson: dict, v: dict) -> dict:
    n = len(lesson["objectives"])
    acts = lesson["activities"]
    execution = [v["exec"].format(i=i + 1, obj=lesson["objectives"][i], pt=lesson["points"][i], act=acts[i % len(acts)]) for i in range(n)]
    return {
        "variant": v["variant"], "name": v["name"], "tagline": v["tagline"],
        "prior": lesson["prior"],
        "objectives": list(lesson["objectives"]),
        "strategies": [{"name": s, "objectives": o} for s, o in v["strategies"](n)],
        "execution": execution,
        "resources": BASE_RESOURCES + lesson["resources"] + v["extra_res"],
        "formative": lesson["formative"] + [v["extra_form"]],
        "enrichment": lesson["enrichment"],
        "remedial": lesson["remedial"],
        "summative": list(lesson["summative"]),
        "homework": v["hw"].format(hw=lesson["homework"]),
        "notes": "",
    }


def lesson_meta(l: dict) -> dict:
    return {"id": l["id"], "grade": l["grade"], "grade_name": GRADE_NAMES[l["grade"]], "unit": l["unit"], "lesson": l["lesson"], "pages": l["pages"]}


PLAN_FIELDS = ("prior", "objectives", "strategies", "execution", "resources", "formative", "enrichment", "remedial", "summative", "homework", "notes")


class PlanSave(BaseModel):
    prior: str = ""
    objectives: List[str] = []
    strategies: List[dict] = []
    execution: List[str] = []
    resources: List[str] = []
    formative: List[str] = []
    enrichment: str = ""
    remedial: str = ""
    summative: List[str] = []
    homework: str = ""
    notes: str = ""


class ExportReq(BaseModel):
    plan: PlanSave
    directorate: Optional[str] = None


class CustomCreate(BaseModel):
    name: Optional[str] = None
    plan: PlanSave


# ============================ DOCX PARSER (استيراد) ============================

GRADE_BY_NAME = {v: k for k, v in GRADE_NAMES.items()}


def _norm(s: str) -> str:
    s = re.sub(r"[\u064B-\u0652\u0640]", "", s or "")
    s = s.replace("أ", "ا").replace("إ", "ا").replace("آ", "ا").replace("ة", "ه").replace("ى", "ي")
    return re.sub(r"\s+", " ", s).strip()


def _row_cells(row):
    out, prev = [], None
    for c in row.cells:
        if c._tc is prev:
            continue
        prev = c._tc
        out.append(c.text)
    return out


def _lines(txt: str):
    return [x.strip() for x in (txt or "").split("\n") if x.strip()]


def _strip_num(s: str):
    return re.sub(r"^\s*[\(\[]?\d+[\)\]\.\-:]\s*", "", s).strip()


def parse_plan_docx(raw: bytes) -> dict:
    doc = Document(BytesIO(raw))
    plan = {k: ([] if k in ("objectives", "strategies", "execution", "resources", "formative", "summative") else "") for k in PLAN_FIELDS}
    meta = {"grade": None, "unit": "", "lesson": ""}
    for table in doc.tables:
        rows = [_row_cells(r) for r in table.rows]
        for i, cells in enumerate(rows):
            joined = " ".join(cells)
            if "الصف" in joined and "الوحدة" in joined and "الدرس" in joined and len(cells) >= 3 and not meta["lesson"]:
                for c in cells:
                    c = c.strip()
                    if c.startswith("الصف"):
                        g = c.split(":", 1)[-1].strip()
                        meta["grade"] = GRADE_BY_NAME.get(g) or next((k for k, v in GRADE_NAMES.items() if v in g), None)
                    elif c.startswith("الوحدة"):
                        meta["unit"] = c.split(":", 1)[-1].strip()
                    elif "الدرس" in c or "الموضوع" in c:
                        meta["lesson"] = re.split(r"[:/]", c, 1)[-1].strip() if (":" in c or "/" in c) else c
                        meta["lesson"] = re.sub(r"^\s*الموضوع\s*[:/]?\s*", "", meta["lesson"]).strip()
            elif "التعلم القبلي" in cells[0] and len(cells) >= 2:
                plan["prior"] = " ".join(_lines(cells[-1]))
            elif "الأهداف" in joined and "الاستراتيجيات" in joined and i + 1 < len(rows):
                body = rows[i + 1]
                if len(body) >= 4:
                    objs = [_strip_num(x) for x in _lines(body[0]) if "يتوقع من الطالب" not in x]
                    plan["objectives"] = [o for o in objs if o]
                    strats = []
                    for ln in _lines(body[1]):
                        m = re.match(r"^\(\s*([^)]*)\)\s*(.+?)\.?$", ln)
                        if m:
                            strats.append({"name": m.group(2).strip(), "objectives": re.sub(r"\s+", "", m.group(1))})
                        else:
                            strats.append({"name": ln.rstrip("."), "objectives": ""})
                    plan["strategies"] = strats
                    plan["execution"] = _lines(body[2])
                    plan["resources"] = [re.sub(r"^[•\-\*]\s*", "", x) for x in _lines(body[3])]
            elif "التقويم التكويني" in joined and "الواجب" in joined and i + 1 < len(rows):
                body = rows[i + 1]
                if len(body) >= 4:
                    plan["formative"] = [re.sub(r"^[•\-\*]\s*", "", x) for x in _lines(body[0])]
                    enr, rem, cur = [], [], None
                    for ln in _lines(body[1]):
                        if "إثرائي" in ln and len(ln) < 25:
                            cur = enr
                        elif "علاجي" in ln and len(ln) < 25:
                            cur = rem
                        elif cur is not None:
                            cur.append(ln)
                        else:
                            enr.append(ln)
                    plan["enrichment"], plan["remedial"] = " ".join(enr), " ".join(rem)
                    plan["summative"] = [re.sub(r"^[•\-\*]\s*", "", x) for x in _lines(body[2])]
                    plan["homework"] = " ".join(_lines(body[3]))
            elif "ملاحظات المعلم" in cells[0]:
                plan["notes"] = " ".join(_lines(cells[-1])) if len(cells) > 1 else ""
    # مطابقة الدرس
    target = _norm(meta["lesson"])
    scored = []
    for l in LESSONS:
        if meta["grade"] and l["grade"] != meta["grade"]:
            continue
        r = SequenceMatcher(None, target, _norm(l["lesson"])).ratio() if target else 0
        if target and (_norm(l["lesson"]) in target or target in _norm(l["lesson"])):
            r = max(r, 0.9)
        scored.append((r, l))
    scored.sort(key=lambda x: -x[0])
    best = scored[0] if scored else (0, None)
    return {
        "detected": {**meta, "grade_name": GRADE_NAMES.get(meta["grade"]) if meta["grade"] else None,
                     "match": lesson_meta(best[1]) if best[1] is not None and best[0] >= 0.6 else None,
                     "confidence": round(best[0], 2),
                     "candidates": [lesson_meta(l) for _, l in scored[:5]]},
        "plan": plan,
        "name": f"مستورد: {meta['lesson']}" if meta["lesson"] else "تحضير مستورد",
    }


def make_router(db, get_teacher):
    router = APIRouter(prefix="/lesson-plans")

    @router.get("/catalog")
    async def catalog(t=Depends(get_teacher)):
        edited = await db.lesson_plan_edits.find({"owner_id": t["teacher_id"]}, {"_id": 0, "lesson_id": 1, "variant": 1}).to_list(1000)
        emap = {}
        for e in edited:
            emap.setdefault(e["lesson_id"], []).append(e["variant"])
        grades = []
        for g in ["5", "6", "7", "8"]:
            units = []
            for l in [x for x in LESSONS if x["grade"] == g]:
                u = next((u for u in units if u["unit"] == l["unit"]), None)
                if not u:
                    u = {"unit": l["unit"], "lessons": []}
                    units.append(u)
                u["lessons"].append({**lesson_meta(l), "edited_variants": sorted(emap.get(l["id"], []))})
            grades.append({"grade": g, "grade_name": GRADE_NAMES[g], "units": units})
        return {"grades": grades, "variants": [{"variant": v["variant"], "name": v["name"], "tagline": v["tagline"]} for v in VARIANTS]}

    async def _customs(owner, lesson_id):
        return await db.lesson_plan_custom.find({"owner_id": owner, "lesson_id": lesson_id}, {"_id": 0}).sort("variant", 1).to_list(50)

    @router.get("/{lesson_id}")
    async def get_plans(lesson_id: str, t=Depends(get_teacher)):
        l = LESSON_INDEX.get(lesson_id)
        if not l:
            raise HTTPException(404, "الدرس غير موجود")
        edits = await db.lesson_plan_edits.find({"owner_id": t["teacher_id"], "lesson_id": lesson_id}, {"_id": 0}).to_list(10)
        emap = {e["variant"]: e["data"] for e in edits}
        variants = []
        for v in VARIANTS:
            base = build_variant(l, v)
            if v["variant"] in emap:
                base = {**base, **{k: emap[v["variant"]].get(k, base[k]) for k in PLAN_FIELDS}, "edited": True}
            else:
                base["edited"] = False
            base["custom"] = False
            variants.append(base)
        for c in await _customs(t["teacher_id"], lesson_id):
            variants.append({"variant": c["variant"], "name": c["name"], "tagline": "تحضير مستورد من ملف Word", "custom": True, "edited": False,
                             **{k: c["data"].get(k, [] if k in ("objectives", "strategies", "execution", "resources", "formative", "summative") else "") for k in PLAN_FIELDS}})
        return {"lesson": lesson_meta(l), "variants": variants}

    @router.post("/import/parse")
    async def import_parse(file: UploadFile = File(...), t=Depends(get_teacher)):
        if not (file.filename or "").lower().endswith(".docx"):
            raise HTTPException(400, "يُقبل ملف Word بصيغة .docx فقط")
        raw = await file.read()
        try:
            parsed = parse_plan_docx(raw)
        except Exception:
            raise HTTPException(400, "تعذر قراءة الملف — تأكد أنه بنفس قالب التحضير")
        if not parsed["plan"]["objectives"] and not parsed["plan"]["execution"]:
            raise HTTPException(400, "لم يتم العثور على جداول التحضير في الملف")
        return parsed

    @router.post("/{lesson_id}/custom")
    async def create_custom(lesson_id: str, req: CustomCreate, t=Depends(get_teacher)):
        if lesson_id not in LESSON_INDEX:
            raise HTTPException(404, "الدرس غير موجود")
        existing = await _customs(t["teacher_id"], lesson_id)
        variant = max([5] + [c["variant"] for c in existing]) + 1
        doc = {"id": str(uuid.uuid4()), "owner_id": t["teacher_id"], "lesson_id": lesson_id, "variant": variant,
               "name": (req.name or "").strip() or f"تحضير مستورد {variant}", "data": req.plan.model_dump(),
               "created_at": datetime.now(timezone.utc).isoformat()}
        await db.lesson_plan_custom.insert_one(doc)
        return {"message": "تمت إضافة التحضير", "variant": variant}

    @router.put("/{lesson_id}/{variant}")
    async def save_plan(lesson_id: str, variant: int, data: PlanSave, t=Depends(get_teacher)):
        if lesson_id not in LESSON_INDEX or variant < 1:
            raise HTTPException(404, "الدرس أو التحضير غير موجود")
        if variant > 5:
            r = await db.lesson_plan_custom.update_one({"owner_id": t["teacher_id"], "lesson_id": lesson_id, "variant": variant},
                                                       {"$set": {"data": data.model_dump(), "updated_at": datetime.now(timezone.utc).isoformat()}})
            if not r.matched_count:
                raise HTTPException(404, "التحضير غير موجود")
            return {"message": "تم حفظ التحضير"}
        await db.lesson_plan_edits.update_one(
            {"owner_id": t["teacher_id"], "lesson_id": lesson_id, "variant": variant},
            {"$set": {"data": data.model_dump(), "updated_at": datetime.now(timezone.utc).isoformat()}}, upsert=True)
        return {"message": "تم حفظ التحضير"}

    @router.delete("/{lesson_id}/{variant}")
    async def reset_plan(lesson_id: str, variant: int, t=Depends(get_teacher)):
        if variant > 5:
            await db.lesson_plan_custom.delete_one({"owner_id": t["teacher_id"], "lesson_id": lesson_id, "variant": variant})
            return {"message": "تم حذف التحضير المستورد"}
        await db.lesson_plan_edits.delete_one({"owner_id": t["teacher_id"], "lesson_id": lesson_id, "variant": variant})
        return {"message": "تمت استعادة التحضير الأصلي"}

    @router.post("/{lesson_id}/{variant}/export")
    async def export_docx(lesson_id: str, variant: int, req: ExportReq, t=Depends(get_teacher)):
        l = LESSON_INDEX.get(lesson_id)
        if not l or variant < 1:
            raise HTTPException(404, "الدرس أو التحضير غير موجود")
        teacher = t["teacher"]
        year = teacher.get("academic_year", "2025-2026").replace("-", "/")
        header = {
            "directorate": req.directorate or teacher.get("directorate") or "المديرية العامة للتربية والتعليم بمحافظة شمال الشرقية",
            "school": teacher.get("school_name") or "مدرسة الخيرات للبنين",
            "teacher": teacher.get("teacher_name") or "",
            "year": year,
        }
        buf = build_docx(l, req.plan.model_dump(), header)
        fname = f"تحضير_{l['lesson']}_{variant}.docx".replace(" ", "_")
        return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                 headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(fname)}"})

    return router


# ============================ DOCX BUILDER ============================

def _rtl_par(p, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p.alignment = align
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)
    return p


def _run(p, text, bold=False, size=11, color=None):
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Arial"
    rPr = r._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), "Arial")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


def _rtl_table(table):
    tblPr = table._tbl.tblPr
    bidi = OxmlElement("w:bidiVisual")
    tblPr.append(bidi)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _cell_lines(cell, lines, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT, label=None):
    cell.text = ""
    first = cell.paragraphs[0]
    if label:
        _rtl_par(first, WD_ALIGN_PARAGRAPH.CENTER)
        _run(first, label, bold=True, size=size)
        first = None
    for ln in lines:
        p = first if first is not None else cell.add_paragraph()
        first = None
        _rtl_par(p, align)
        _run(p, ln, bold=bold, size=size)


def build_docx(lesson: dict, plan: dict, header: dict) -> BytesIO:
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.left_margin = sec.right_margin = Cm(1.2)
    sec.top_margin = sec.bottom_margin = Cm(1.0)

    for text, size, bold in [(header["directorate"], 13, True), (f"مدرسة: {header['school']}", 12, True),
                             (f"تحضير مادة تقنية المعلومات   العام الدراسي {header['year']} م", 12, True), (f"اسم المعلم/ {header['teacher']}", 11, False)]:
        p = _rtl_par(doc.add_paragraph(), WD_ALIGN_PARAGRAPH.CENTER)
        _run(p, text, bold=bold, size=size)

    # جدول الصف/الوحدة/الدرس
    t0 = doc.add_table(rows=1, cols=3)
    _rtl_table(t0)
    for c, txt in zip(t0.rows[0].cells, [f"الصف: {GRADE_NAMES[lesson['grade']]}", f"الوحدة: {lesson['unit']}", f"عنوان الدرس/ الموضوع: {lesson['lesson']}  (ص {lesson['pages']})"]):
        _cell_lines(c, [txt], bold=True, size=11)
        _shade(c, "E7E6E6")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # جدول التعلم القبلي + الأهداف/الاستراتيجيات/التنفيذ/الوسائل
    t1 = doc.add_table(rows=3, cols=4)
    _rtl_table(t1)
    r0 = t1.rows[0]
    m = r0.cells[1].merge(r0.cells[3])
    _cell_lines(r0.cells[0], ["التعلم القبلي/التمهيد/ المفاهيم"], bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade(r0.cells[0], "E7E6E6")
    _cell_lines(m, [plan["prior"]], size=10)
    heads = ["الأهداف/ المخرجات التعليمية", "الاستراتيجيات/طرق التدريس", "آلية التنفيذ/ الأنشطة التدريبية/التعليمية", "الوسائل ومصادر التعلم"]
    for c, h in zip(t1.rows[1].cells, heads):
        _cell_lines(c, [h], bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(c, "D9E2F3")
    r2 = t1.rows[2].cells
    _cell_lines(r2[0], ["يتوقع من الطالب بعد انتهاء الدرس أن :"] + [f"{i + 1}. {o}" for i, o in enumerate(plan["objectives"])], size=10)
    _cell_lines(r2[1], [f"(  {s.get('objectives', '')}  ) {s.get('name', '')}." for s in plan["strategies"]], size=10)
    _cell_lines(r2[2], list(plan["execution"]), size=10)
    _cell_lines(r2[3], [f"• {r}" for r in plan["resources"]], size=10)
    widths = [Cm(6.2), Cm(5.0), Cm(11.3), Cm(4.8)]
    for row in t1.rows[1:]:
        for c, w in zip(row.cells, widths):
            c.width = w

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # جدول التقويم
    t2 = doc.add_table(rows=3, cols=4)
    _rtl_table(t2)
    heads2 = ["التقويم التكويني", "نشاط إثرائي/ علاجي / تفريد التعليم", "التقويم الختامي", "الواجب المنزلي"]
    for c, h in zip(t2.rows[0].cells, heads2):
        _cell_lines(c, [h], bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(c, "D9E2F3")
    rr = t2.rows[1].cells
    _cell_lines(rr[0], [f"- {q}" for q in plan["formative"]], size=10)
    _cell_lines(rr[1], ["نشاط إثرائي:", plan["enrichment"], "نشاط علاجي:", plan["remedial"]], size=10)
    _cell_lines(rr[2], [f"- {q}" for q in plan["summative"]], size=10)
    _cell_lines(rr[3], [plan["homework"]], size=10)
    last = t2.rows[2]
    mm = last.cells[1].merge(last.cells[3])
    _cell_lines(last.cells[0], ["ملاحظات المعلم"], bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade(last.cells[0], "E7E6E6")
    _cell_lines(mm, [plan.get("notes") or " "], size=10)

    p = _rtl_par(doc.add_paragraph(), WD_ALIGN_PARAGRAPH.CENTER)
    _run(p, "يعتمد،، المعلم الأول" + " " * 90 + "يعتمد،،، المشرف التربوي", bold=True, size=11)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
